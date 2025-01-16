import os
from docx import Document

def generate_docx_from_txt(txt_dir, docx_dir):
    """Generate DOCX files from TXT files in the specified directory."""
    try:
        # Ensure the DOCX output directory exists
        os.makedirs(docx_dir, exist_ok=True)

        # List all TXT files in the directory
        txt_files = [f for f in os.listdir(txt_dir) if f.endswith('.txt')]

        for txt_file in txt_files:
            txt_path = os.path.join(txt_dir, txt_file)
            base_filename = os.path.splitext(txt_file)[0]

            # Read the content of the TXT file
            with open(txt_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Create a new DOCX file
            docx_output_path = os.path.join(docx_dir, f"{base_filename}.docx")
            doc = Document()
            doc.add_heading('Content', level=1)
            doc.add_paragraph(content)
            doc.save(docx_output_path)

            print(f"DOCX file created: {docx_output_path}")

    except Exception as e:
        print(f"Error generating DOCX files: {e}")

if __name__ == "__main__":
    txt_directory = "results/2025-01-16/txt"
    docx_directory = "results/2025-01-16/docx"
    generate_docx_from_txt(txt_directory, docx_directory)
