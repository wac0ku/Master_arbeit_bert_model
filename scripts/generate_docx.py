from docx import Document
from docx.shared import Cm
from datetime import datetime
from pathlib import Path
import pandas as pd
import matplotlib.pyplot as plt
import io

def create_analysis_report(results: dict, output_dir: Path) -> str:
    """Generiere Masterarbeit-Analysebericht im DOCX-Format"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_filename = f"Masterarbeit_Analysebericht_{timestamp}.docx"
    report_path = output_dir / report_filename

    doc = Document()
    
    # Titelblatt
    doc.add_heading('Masterarbeit - Sicherheitsanalysebericht', 0)
    doc.add_paragraph(f"Generiert am: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
    
    # Statistische Auswertung
    doc.add_heading('Gesamtstatistik', level=1)
    risk_data = pd.DataFrame.from_dict(results, orient='index')
    
    # Risikoverteilung
    plt.figure(figsize=(10, 6))
    risk_data['risk_level'].value_counts().plot.bar()
    plt.title('Verteilung der Risikobewertungen')
    plt.xlabel('Risikostufe')
    plt.ylabel('Anzahl')
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150)
    doc.add_picture(img_stream, width=Cm(15))
    
    # Detailtabelle
    doc.add_heading('Einzelergebnisse', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = 'Dateiname'
    hdr[1].text = 'Risikostufe'
    hdr[2].text = 'Empfehlungen'
    hdr[3].text = 'Schlüsselwörter'
    
    # Datenzeilen
    for filename, data in results.items():
        row = table.add_row().cells
        row[0].text = filename
        row[1].text = data['risk_level'].upper()
        row[2].text = '\n'.join(data['recommendations'])
        row[3].text = str(data['keyword_count'])
    
    # Speichern
    doc.save(report_path)
    return str(report_path)