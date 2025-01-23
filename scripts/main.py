# main.py
from rich.console import Console
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, BarColumn, TimeRemainingColumn
from pathlib import Path
from datetime import datetime
import json
import pdfplumber
from docx import Document
from modules.causality import analyze_easa_causality
from typing import Dict, Any

console = Console()

# Pfadkonfiguration
BASE_DIR = Path(__file__).parent.parent
DATA_DIR = BASE_DIR / "data"
RAW_PDF_DIR = DATA_DIR / "raw_pdf"
TEMPLATES_DIR = BASE_DIR / "easa_templates"
CATEGORIES_PATH = TEMPLATES_DIR / "categories.json"

class EASAProcessor:
    def __init__(self):
        with open(CATEGORIES_PATH, 'r') as f:
            self.categories = json.load(f)['easa_risk_categories']
            self.compliance_matrix = json.load(f)['compliance_matrix']

    def process_pdf(self, pdf_path: Path, output_dir: Path) -> Dict[str, Any]:
        """Verarbeite PDF gemäß EASA VO (EU) 376/2014"""
        result = {
            'filename': pdf_path.name,
            'timestamp': datetime.now().isoformat(),
            'categories': [],
            'recommendations': []
        }

        try:
            with pdfplumber.open(pdf_path) as pdf:
                full_text = "\n".join(
                    f"Page {p.page_number}:\n{p.extract_text()}\n" 
                    for p in pdf.pages
                )

                # Analyse
                result.update({
                    'causal_chains': analyze_easa_causality(full_text),
                    'categories': self._detect_categories(full_text),
                    'recommendations': self._generate_recommendations(full_text)
                })

                # Dateien speichern
                self._save_reports(result, full_text, output_dir)

        except Exception as e:
            console.print(f"[red]Fehler bei {pdf_path.name}: {str(e)}")
            result['error'] = str(e)

        return result

    def _detect_categories(self, text: str) -> list:
        """Erkenne Risikokategorien gemäß categories.json"""
        detected = []
        text_lower = text.lower()
        for cat in self.categories:
            if any(ex.lower() in text_lower for ex in cat['examples']):
                detected.append({
                    'id': cat['id'],
                    'name': cat['name'],
                    'priority': cat['priority']
                })
        return detected

    def _generate_recommendations(self, text: str) -> list:
        """Generiere EASA-konforme Empfehlungen"""
        recs = []
        text_lower = text.lower()
        
        for cat in self.categories:
            if any(ex.lower() in text_lower for ex in cat['examples']):
                recs.append({
                    'category': cat['id'],
                    'priority': cat['priority'],
                    'action': f"{cat['easa_reference']}: {cat['examples'][0]}",
                    'deadline': self.compliance_matrix['priority_levels'][cat['priority']]
                })
        return recs[:4]

    def _save_reports(self, data: dict, text: str, output_dir: Path):
        """Speichere Berichte im EASA-Format"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # TXT Report
        txt_dir = output_dir / "txt"
        txt_dir.mkdir(exist_ok=True)
        with open(txt_dir / f"{data['filename']}_analysis_{timestamp}.txt", 'w') as f:
            f.write(f"EASA Safety Analysis Report\n{'='*40}\n")
            f.write(f"File: {data['filename']}\nDate: {data['timestamp']}\n\n")
            f.write("Causal Chains:\n" + "\n".join(data['causal_chains']) + "\n\n")
            f.write("Recommendations:\n" + "\n".join([r['action'] for r in data['recommendations']]))

        # DOCX Report
        docx_dir = output_dir / "docx"
        docx_dir.mkdir(exist_ok=True)
        doc = Document()
        
        # Header
        doc.add_heading('EASA Safety Analysis Report', 0)
        doc.add_paragraph(f"Generated: {data['timestamp']}")
        
        # Hauptabschnitte
        self._add_summary_section(doc, text)
        self._add_causal_analysis(doc, data['causal_chains'])
        self._add_recommendations(doc, data['recommendations'])
        
        doc.save(docx_dir / f"{data['filename']}_analysis_{timestamp}.docx")

    def _add_summary_section(self, doc: Document, text: str):
        """Füge Zusammenfassung hinzu"""
        doc.add_heading('Executive Summary', 1)
        summary = text[:500] + "[...]" if len(text) > 500 else text
        doc.add_paragraph(summary)
        doc.add_paragraph(f"Total length: {len(text)} characters")

    def _add_causal_analysis(self, doc: Document, chains: list):
        """Füge Fehlerkettenanalyse hinzu"""
        doc.add_heading('Causal Chain Analysis (CCA)', 1)
        for chain in chains:
            doc.add_paragraph(chain, style='ListBullet')
        if not chains:
            doc.add_paragraph("No causal chains detected", style='Intense Quote')

    def _add_recommendations(self, doc: Document, recommendations: list):
        """Füge Empfehlungen hinzu"""
        doc.add_heading('Safety Recommendations', 1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header
        hdr = table.rows[0].cells
        hdr[0].text = 'Priority'
        hdr[1].text = 'Category'
        hdr[2].text = 'Required Action'
        hdr[3].text = 'Deadline'

        # Daten
        for rec in recommendations:
            row = table.add_row().cells
            row[0].text = rec['priority']
            row[1].text = rec['category']
            row[2].text = rec['action']
            row[3].text = rec['deadline']

def main():
    console.print(Panel.fit(
        "[bold blue]EASA Safety Analysis System v3.0\n"
        "EU Regulation 376/2014 Compliant",
        subtitle=f"Generated {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    ))

    try:
        # Verzeichnisstruktur
        analysis_date = datetime.now().strftime("%Y-%m-%d")
        output_dir = BASE_DIR / "analysis_results" / analysis_date
        output_dir.mkdir(parents=True, exist_ok=True)

        # Initialisierung
        processor = EASAProcessor()
        pdf_files = list(RAW_PDF_DIR.glob("*.pdf"))

        # Verarbeitung
        with Progress(
            SpinnerColumn(),
            "[progress.description]{task.description}",
            BarColumn(bar_width=40),
            TimeRemainingColumn()
        ) as progress:
            task = progress.add_task("[cyan]Processing PDFs...", total=len(pdf_files))

            for pdf_file in pdf_files:
                result = processor.process_pdf(pdf_file, output_dir)
                progress.update(task, advance=1, 
                    description=f"Processing {pdf_file.name[:15]}...")

        console.print(Panel.fit(
            f"[bold green]✓ Analysis completed successfully!\n"
            f"Reports saved to: [link file://{output_dir}]{output_dir}",
            border_style="green"
        ))

    except Exception as e:
        console.print(Panel.fit(
            f"[bold red]✗ Critical system failure\n{str(e)}",
            border_style="red"
        ))
        raise

if __name__ == "__main__":
    main()