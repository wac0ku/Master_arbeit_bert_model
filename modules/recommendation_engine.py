import json
from pathlib import Path
from docx import Document
from datetime import datetime

class EASARecommendationGenerator:
    def __init__(self, categories_path: str):
        with open(categories_path, 'r') as f:
            data = json.load(f)
            self.categories = data['easa_risk_categories']
            self.compliance_matrix = data['compliance_matrix']
    
    def _match_category(self, text: str) -> dict:
        detected = []
        text_lower = text.lower()
        for cat in self.categories:
            if any(ex.lower() in text_lower for ex in cat['examples']):
                detected.append(cat)
        return detected
    
    def generate_recommendations(self, analysis_text: str) -> dict:
        detected_cats = self._match_category(analysis_text)
        recommendations = []
        
        for cat in detected_cats:
            recommendation = {
                'category_id': cat['id'],
                'priority': cat['priority'],
                'action_deadline': self.compliance_matrix['priority_levels'][cat['priority']],
                'required_actions': [
                    f"Implement {cat['name']} mitigation plan",
                    f"Review {cat['examples'][0]} procedures",
                    f"Submit report to EASA per {cat['easa_reference']}"
                ]
            }
            recommendations.append(recommendation)
        
        return {
            'detected_categories': detected_cats,
            'recommendations': recommendations,
            'reporting_requirements': self.compliance_matrix['reporting_requirements']['mandatory_reporting']
        }

    def create_docx_report(self, analysis_text: str, output_path: Path):
        doc = Document()
        analysis_data = self.generate_recommendations(analysis_text)
        
        # Header
        doc.add_heading('EASA Safety Analysis Report', 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        # Summary Section
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(analysis_text[:500] + "...")  # First 500 chars as summary
        
        # Detailed Analysis
        doc.add_heading('Causal Chain Analysis', 1)
        for category in analysis_data['detected_categories']:
            doc.add_heading(f"Category {category['id']}: {category['name']}", 2)
            doc.add_paragraph(category['description'])
            doc.add_paragraph(f"**Regulatory Reference:** {category['easa_reference']}")
        
        # Recommendations
        doc.add_heading('Safety Recommendations', 1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Priority'
        hdr_cells[1].text = 'Category'
        hdr_cells[2].text = 'Action'
        hdr_cells[3].text = 'Deadline'
        
        for rec in analysis_data['recommendations']:
            row_cells = table.add_row().cells
            row_cells[0].text = rec['priority']
            row_cells[1].text = rec['category_id']
            row_cells[2].text = "\n".join(rec['required_actions'])
            row_cells[3].text = rec['action_deadline']
        
        doc.save(output_path)